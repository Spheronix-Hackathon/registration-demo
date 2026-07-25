from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO

# Fix M-10: Use IST-aware timestamp from app time_utils instead of datetime.datetime.now()
from app.core.time_utils import format_ist

def generate_problem_statements_docx():
    doc = Document()
    
    # Title
    title = doc.add_heading('Official Problem Statements', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_paragraph('Spheronix Hackathon 2026')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.color.rgb = RGBColor(0, 122, 255) # iOS Accent Blue
    
    doc.add_paragraph(f"Generated on: {format_ist()}")
    doc.add_paragraph("-" * 50)
    
    # Track 01: Full Stack
    doc.add_heading('Track 01: Complete Full Stack Application', level=1)
    
    doc.add_heading('Problem Statement', level=2)
    doc.add_paragraph(
        "Organizations and startups require efficient web-based systems to manage users, data, and services. "
        "Building an integrated application combining frontend, backend, security, and database management is complex. "
        "Many existing solutions lack scalability and security."
    )
    
    doc.add_heading('Objectives', level=2)
    objectives = [
        "Provide robust user authentication and authorization.",
        "Enable efficient data management and retrieval.",
        "Ensure smooth frontend and backend communication.",
        "Deliver a responsive and user-friendly interface.",
        "Adaptable to multiple real-world applications."
    ]
    for obj in objectives:
        doc.add_paragraph(obj, style='List Bullet')
        
    doc.add_heading('System Architecture', level=2)
    doc.add_paragraph("1. Frontend Layer: Provides interface for user interaction.", style='List Bullet')
    doc.add_paragraph("2. Backend Layer: Handles processing and management logic.", style='List Bullet')
    doc.add_paragraph("3. Database Layer: Stores user data and ensures consistency.", style='List Bullet')
    
    doc.add_heading('Technologies Used', level=2)
    doc.add_paragraph("React.js, Node.js/Express, Python/Java, MongoDB/MySQL, JWT Auth, REST APIs")
    
    doc.add_heading('Key Features', level=2)
    features = [
        "Secure user registration and login system.",
        "Role-based or protected access permissions.",
        "Dynamic dashboard for data visualization.",
        "Responsive and mobile-friendly UI layout."
    ]
    for feat in features:
        doc.add_paragraph(feat, style='List Bullet')
        
    doc.add_page_break()
    
    # Track 02: Bug Hunting
    doc.add_heading('Track 02: Bug Hunting Web Application', level=1)
    
    doc.add_heading('Problem Statement', level=2)
    doc.add_paragraph(
        "Undetected bugs lead to data breaches, system failures, and financial loss. Traditional debugging is manual and inefficient. "
        "Many solutions lack a unified approach across the entire application stack."
    )
    
    doc.add_heading('Objectives', level=2)
    objectives2 = [
        "Identify bugs effectively in web applications.",
        "Track and manage reported issues efficiently.",
        "Improve application security and performance.",
        "Provide a user-friendly interface for monitoring.",
        "Enhance overall system reliability."
    ]
    for obj in objectives2:
        doc.add_paragraph(obj, style='List Bullet')
        
    doc.add_heading('Architecture Flow', level=2)
    flow = [
        "User identifies bug and submits report with evidence.",
        "Frontend sends data to backend via APIs.",
        "Backend processes and stores information in Database.",
        "Developers access, track, and update bug status.",
        "System sends real-time updates and notifications."
    ]
    for item in flow:
        doc.add_paragraph(item, style='List Number')
        
    doc.add_heading('Key Features', level=2)
    features2 = [
        "Detailed bug reporting with screenshots/metadata.",
        "Categorization (UI, Functional, Security, Performance).",
        "Comprehensive status tracking (Open, Resolved).",
        "Secure user and developer authentication.",
        "Real-time updates and notification system."
    ]
    for feat in features2:
        doc.add_paragraph(feat, style='List Bullet')
        
    doc.add_page_break()
    
    # Track 03: Native Windows
    doc.add_heading('Track 03: Native Windows Application', level=1)
    
    doc.add_heading('Problem Statement', level=2)
    doc.add_paragraph(
        "Web applications often face connectivity dependencies and reduced performance. "
        "Windows users requires high-performance applications that fully utilize system capabilities, "
        "hardware integration, and offline functionality."
    )
    
    doc.add_heading('Objectives', level=2)
    objectives3 = [
        "Provide high performance and responsiveness.",
        "Work efficiently in offline and online modes.",
        "Utilize system-level features and hardware.",
        "Offer a user-friendly and intuitive native interface.",
        "Adaptable for real-world productivity tools."
    ]
    for obj in objectives3:
        doc.add_paragraph(obj, style='List Bullet')
        
    doc.add_heading('Native Tech Stack', level=2)
    doc.add_paragraph("C# / WinForms / WPF, .NET Core / .NET Framework, SQLite / Local Storage, Windows SDK, Visual Studio")
    
    doc.add_heading('Expected Outcome', level=2)
    outcomes = [
        "A fully functional native Windows application.",
        "Superior performance compared to web equivalents.",
        "Smooth user experience with offline capabilities.",
        "Scalable solution for desktop environments."
    ]
    for outcome in outcomes:
        doc.add_paragraph(outcome, style='List Bullet')
        
    # Save to buffer
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream
